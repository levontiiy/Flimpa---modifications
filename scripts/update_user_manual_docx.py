"""Update User Manual FLIMPA v2.0.0.docx with new TOC and 2.0.0 sections."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

DOC_PATH = Path("/Users/levgerasimov/Downloads/User Manual FLIMPA v2.0.0.docx")
BACKUP_PATH = DOC_PATH.with_name("User Manual FLIMPA v2.0.0_backup.docx")

TOC_LINES = [
    "1. Installation and getting started",
    "    1.1 System requirements",
    "    1.2 Installing from release — macOS (.dmg)",
    "    1.3 Installing from release — Windows (.exe)",
    "    1.4 Installing from source",
    "    1.5 First launch and main window layout",
    "    1.6 Menus: Load data, Reference, Mask save, Save data",
    "2. Loading data",
    "    2.1 Supported file formats",
    "    2.2 Load data menu options",
    "    2.3 Loading .tif and .tiff files",
    "    2.4 Loading .ptu files",
    "    2.5 Loading data with manual masks",
    "    2.6 File list and sample data",
    "3. Reference and calibration",
    "    3.1 Import reference file",
    "    3.2 Reference lifetime and donor τ_D for FRET",
    "    3.3 Import IRF",
    "4. Running phasor plot analysis",
    "5. Phasor plot",
    "    5.1 G and S axis labels",
    "    5.2 Visualisation modes and gallery Layers list",
    "    5.3 Phasor ROI (ellipse)",
    "6. Image tabs and navigation",
    "    6.1 Intensity display and colormaps",
    "    6.2 Lifetime maps",
    "    6.3 FRET efficiency maps",
    "    6.4 Pan, zoom, and reset",
    "7. Baseline check (decay curve inspector)",
    "8. Masking",
    "    8.1 Masking tools (polygon, lasso, brush, delete region, eraser)",
    "    8.2 Mask save menu",
    "    8.3 Importing and re-importing saved masks",
    "9. Results overview",
    "    9.1 Lifetime values table",
    "    9.2 Lifetime maps",
    "    9.3 Gallery (tau) and Gallery (I)",
    "    9.4 Violin plots",
    "10. Saving and exporting data",
    "    10.1 Save data menu",
    "    10.2 Export phasor points (G, S)",
    "11. Understanding the outputs (G, S, M, phi, row, col)",
    "12. Troubleshooting and FAQ",
    "Appendix A. Changes from FLIMPA 1.4.2 to 2.0.0",
]

GETTING_STARTED = {
    "heading": "1. Installation and getting started",
    "subsections": [
        (
            "1.1 System requirements",
            "Python 3.11+ for install from source. macOS, Windows, or Linux. "
            "For best results use images up to 512×512 pixels; larger images may be slow or run out of memory.",
        ),
        (
            "1.2 Installing from release — macOS (.dmg)",
            "Download FLIMPA.v2.0.0.dmg from GitHub Releases (levontiiy/Flimpa---modifications). "
            "Open the disk image and drag FLIMPA to Applications. "
            "If macOS blocks the app on first launch (unsigned build), right-click the app and choose Open.",
        ),
        (
            "1.3 Installing from release — Windows (.exe)",
            "[Content to be added when the Windows release is published on GitHub Releases.]",
        ),
        (
            "1.4 Installing from source",
            "Clone or download the repository, create a virtual environment, run "
            "pip install -r requirements.txt, then python main.py. "
            "See the project README on GitHub for conda/venv instructions.",
        ),
        (
            "1.5 First launch and main window layout",
            "Parameters and Run Phasor Plot Analysis are on the left; the phasor plot is below that; "
            "image tabs are in the centre; the file list is on the right. "
            "After analysis, tabs appear for lifetime maps, galleries, violin plots, and the lifetime table.",
        ),
        (
            "1.6 Menus",
            "Load data — import samples and masks. Reference — reference file and IRF. "
            "Mask save — save or clear manual and ROI masks. Save data — export maps, galleries, plots, and CSV files.",
        ),
    ],
    "sample_data": (
        "Sample Becker & Hickl .sdt files and example mask TIFFs are in the project sample_data folder on GitHub."
    ),
}


def _clear_paragraph(p):
    p.text = ""


def _set_paragraph_text(p, text, *, bold=False):
    _clear_paragraph(p)
    run = p.add_run(text)
    if bold:
        run.bold = True


def _insert_toc_after_contents(doc: Document) -> None:
    idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Contents")
    insert_at = idx + 1
    # Remove old TOC / blank lines until "User Manual"
    while insert_at < len(doc.paragraphs):
        t = doc.paragraphs[insert_at].text.strip()
        if t == "User Manual":
            break
        _clear_paragraph(doc.paragraphs[insert_at])
        insert_at += 1
    # Insert new TOC lines before "User Manual"
    user_manual_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "User Manual")
    anchor = doc.paragraphs[user_manual_idx]
    for line in TOC_LINES:
        new_p = anchor.insert_paragraph_before(line)
        new_p.style = doc.styles["Normal"]
        new_p.paragraph_format.space_after = Pt(0)


def _replace_getting_started(doc: Document) -> None:
    idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Getting started")
    doc.paragraphs[idx].text = GETTING_STARTED["heading"]
    doc.paragraphs[idx].style = doc.styles["Heading 1"]

    # Clear old list paragraphs until "Loading data"
    end = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Loading data")
    for j in range(idx + 1, end):
        _clear_paragraph(doc.paragraphs[j])

    blocks: list[tuple[str, str]] = []
    for title, body in GETTING_STARTED["subsections"]:
        blocks.append((title, "Heading 2"))
        blocks.append((body, "List Paragraph"))
    blocks.append((GETTING_STARTED["sample_data"], "List Paragraph"))

    anchor = doc.paragraphs[end]
    for text, style_name in reversed(blocks):
        new_p = anchor.insert_paragraph_before(text)
        new_p.style = doc.styles[style_name]


def _update_saving_section(doc: Document) -> None:
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "Saving FLIMPA’s outputs" or t == "Saving FLIMPA's outputs":
            p.text = "10. Saving and exporting data"
        if "Save" in t and "button" in t and "toolbar" in t:
            p.text = (
                "Figure 14. Data can be exported from the Save data menu in the top menu bar. "
                "Masks are saved separately from the Mask save menu."
            )
        if t == "FLIMPA allows users to export all generated data. This includes:":
            p.text = (
                "The Save data menu allows export of all generated data. This includes:"
            )

    # Append new export bullets at end if not present
    last_items = [p.text.strip() for p in doc.paragraphs[-8:] if p.text.strip()]
    if not any("phasor points" in t.lower() for t in last_items):
        doc.add_paragraph(
            "Export phasor points (G,S): choose an analysed file, then save a CSV with columns "
            "G, S, row, col (non-zero pixels only).",
            style="List Paragraph",
        )
        doc.add_paragraph(
            "Export lifetime values table: CSV of mean lifetime per image.",
            style="List Paragraph",
        )


def _add_new_sections(doc: Document) -> None:
    if any(p.text.strip().startswith("5. Phasor plot (2.0.0 updates)") for p in doc.paragraphs):
        return

    sections = [
        (
            "5. Phasor plot (2.0.0 updates)",
            "[Add: G/S axis labels, Layers list on gallery phasor plots, scatter/histogram/contour, Individual vs Condition.]",
        ),
        (
            "6. Image tabs and navigation",
            "[Add: Intensity colormaps, FRET tab (rightmost after analysis), pan/zoom/reset on image views.]",
        ),
        (
            "7. Baseline check (decay curve inspector)",
            "[Add: enable on Lifetime maps tab, click pixel, log scale, map τ curve, IRF overlay, t₀ crop, Move map τ.]",
        ),
        (
            "8. Masking",
            "[Add: Masking tools, eraser/brush size, Mask save menu, mask re-import after save, labelled uint16 TIFF format.]",
        ),
        (
            "11. Understanding the outputs",
            "[Add: relationship between G, S, M, phi, image row/col, and FRET efficiency.]",
        ),
        (
            "12. Troubleshooting and FAQ",
            "[Add: masks not visible until re-import, macOS Gatekeeper, baseline correction warnings.]",
        ),
        (
            "Appendix A. Changes from FLIMPA 1.4.2 to 2.0.0",
            "[Summary of new features in this fork.]",
        ),
    ]
    for title, body in sections:
        doc.add_paragraph(title, style="Heading 1")
        doc.add_paragraph(body, style="Normal")


def main():
    if not DOC_PATH.exists():
        raise SystemExit(f"Not found: {DOC_PATH}")

    shutil.copy2(DOC_PATH, BACKUP_PATH)
    doc = Document(str(DOC_PATH))
    _insert_toc_after_contents(doc)
    _replace_getting_started(doc)
    _update_saving_section(doc)
    _add_new_sections(doc)
    doc.save(str(DOC_PATH))
    print(f"Updated: {DOC_PATH}")
    print(f"Backup:  {BACKUP_PATH}")


if __name__ == "__main__":
    main()
